"""Shared helpers for Word (.docx) report generation."""

from __future__ import annotations

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def shade_header_row(table) -> None:
    for cell in table.rows[0].cells:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "1F4E79")
        cell._tc.get_or_add_tcPr().append(shading)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, label in enumerate(headers):
        table.rows[0].cells[i].text = label
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, value in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = value
    shade_header_row(table)


def add_body(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15


def add_code(doc: Document, text: str) -> None:
    for line in text.splitlines():
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0


def page_break(doc: Document) -> None:
    doc.add_page_break()
