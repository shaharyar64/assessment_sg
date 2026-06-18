"""Shared Word document sections for the final report."""

from __future__ import annotations

from typing import Any

from docx import Document

from src.docx_helpers import add_body, add_code, add_table
from src.evaluator import METRIC_DEFINITIONS
from src.prompt_builder import (
    STRATEGY_BASIC,
    STRATEGY_STRUCTURED,
    build_basic_prompt,
    build_structured_prompt,
)


def write_prompts_section(
    doc: Document,
    sample: dict[str, Any] | None,
    winner: str,
) -> None:
    doc.add_heading("1. Prompt Templates Used", level=1)
    add_body(
        doc,
        f"Two LangChain ChatPromptTemplate strategies were evaluated. "
        f"Production recommendation: {winner}.",
    )

    doc.add_heading("Strategy A — Basic", level=2)
    add_body(doc, "Minimal template: system message, intent, key facts, tone, four sections.")
    if sample:
        add_body(doc, f"Example (scenario {sample.get('id', 'SCN-001')}):")
        add_code(doc, build_basic_prompt(sample))

    doc.add_heading("Strategy B — Structured", level=2)
    add_body(
        doc,
        "Role-based template with tone guide, structure rules, no-hallucination constraints, "
        "and explicit output format.",
    )
    if sample:
        add_body(doc, f"Example (scenario {sample.get('id', 'SCN-001')}):")
        add_code(doc, build_structured_prompt(sample))


def write_metrics_section(doc: Document) -> None:
    doc.add_heading("2. Custom Evaluation Metrics — Definitions and Logic", level=1)
    add_body(
        doc,
        "Each generated email is scored on three metrics (0.0–1.0). "
        "Overall average = mean of the three scores.",
    )
    for idx, meta in enumerate(METRIC_DEFINITIONS.values(), start=1):
        doc.add_heading(f"2.{idx} {meta['name']}", level=2)
        add_body(doc, f"Definition: {meta['definition']}")
        add_body(doc, f"Logic: {meta['logic']}")


def write_comparative_analysis_section(
    doc: Document,
    comparison: dict[str, Any],
) -> None:
    winner = comparison.get("recommended_strategy", comparison.get("winner", "N/A"))
    loser = comparison.get("loser", "N/A")
    avgs = comparison.get("overall_averages", {})
    metric_comp = comparison.get("metric_comparison", {})

    doc.add_heading("3. Comparative Analysis Summary", level=1)
    add_body(
        doc,
        f"{STRATEGY_BASIC} (Strategy A) vs {STRATEGY_STRUCTURED} (Strategy B) on 10 scenarios.",
    )

    doc.add_heading("Which strategy performed better?", level=2)
    add_body(doc, f"Winner: {winner}")
    add_body(doc, comparison.get("reason", ""))

    doc.add_heading("Overall averages", level=2)
    add_table(
        doc,
        ["Strategy", "Overall Average"],
        [[name, f"{avgs[name]:.4f}"] for name in sorted(avgs, key=lambda s: avgs[s], reverse=True)],
    )

    doc.add_heading("Metric-by-metric comparison", level=2)
    add_table(
        doc,
        ["Metric", winner, loser],
        [
            [name, f"{scores.get(winner, 0):.4f}", f"{scores.get(loser, 0):.4f}"]
            for name, scores in metric_comp.items()
        ],
    )

    doc.add_heading("Biggest failure mode (lower-performing strategy)", level=2)
    add_body(doc, f"Strategy: {loser}")
    add_body(doc, comparison.get("biggest_failure_mode", "N/A"))

    doc.add_heading("Production recommendation", level=2)
    add_body(doc, comparison.get("production_recommendation", ""))


def write_raw_data_section(doc: Document, evaluation_results: dict[str, Any]) -> None:
    doc.add_heading("4. Raw Evaluation Data", level=1)
    add_body(
        doc,
        "Per-scenario scores below match outputs/evaluation_results.json and "
        "outputs/evaluation_results.csv (20 rows: 10 scenarios × 2 strategies).",
    )

    add_table(
        doc,
        ["Scenario", "Strategy", "Fact Recall", "Tone & Format", "Conciseness", "Overall"],
        [
            [
                row.get("scenario_id", ""),
                row.get("strategy", ""),
                f"{row.get('scores', {}).get('fact_recall', 0):.3f}",
                f"{row.get('scores', {}).get('tone_and_format_accuracy', 0):.3f}",
                f"{row.get('scores', {}).get('conciseness_and_fluency', 0):.3f}",
                f"{row.get('overall_average', 0):.3f}",
            ]
            for row in evaluation_results.get("scenario_scores", [])
        ],
    )

    doc.add_heading("Strategy averages", level=2)
    add_table(
        doc,
        ["Strategy", "Fact Recall", "Tone & Format", "Conciseness", "Overall"],
        [
            [
                strategy,
                f"{stats.get('metric_averages', {}).get('fact_recall', 0):.3f}",
                f"{stats.get('metric_averages', {}).get('tone_and_format_accuracy', 0):.3f}",
                f"{stats.get('metric_averages', {}).get('conciseness_and_fluency', 0):.3f}",
                f"{stats.get('overall_average', 0):.3f}",
            ]
            for strategy, stats in evaluation_results.get("strategy_averages", {}).items()
        ],
    )
